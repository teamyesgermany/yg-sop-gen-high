import streamlit as st
import io
import os
import openai as ai
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import spacy
import random



load_dotenv()  # take environment variables from .env.

ai.api_key = os.getenv("OPENAI_API_KEY")

nlp = spacy.blank("en")  # Create a blank English model
#^ Load the SpaCy English model
nlp_e = spacy.load("en_core_web_sm")


def retrieve_cgpa_score(content_resume):
    
    # Check if the entity_ruler already exists in the pipeline
    if 'entity_ruler' in nlp.pipe_names:
        ruler = nlp.get_pipe('entity_ruler')
    else:
        ruler = nlp.add_pipe("entity_ruler")

    patterns = [
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.6 CGPA"
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "CGPA 9.6" and "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]},  # Matches "CGPA - 9.26" or "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "9.5 CGPA" or "9.5 CGPA -" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.5 - CGPA" or "9.5 . CGPA" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["-", "–"]}}, {"IS_SPACE": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]}, # Matches "CGPA - 9.26" or "CGPA -9.5" (with optional punctuation and space)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\w*\\d+(\\.\\d+)?"}}, {"LOWER": {"REGEX": "(cgpa|sgpi\\w*)"}}]},
        {"label": "CGPA", "pattern": [{"LOWER": {"REGEX": "(\\w*cgpa|sgpi)"}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?\w*"}}]}
    ]

    ruler.add_patterns(patterns)

    lines = content_resume.split('\n')
    score = "unknown"
    for index, line in enumerate(lines):
        doc = nlp(line)
        found_cgpa = False
        for ent in doc.ents:
            if ent.label_ == 'CGPA':
                found_cgpa = True
                print(ent.text)
                numerical_score = nlp_e(ent.text)
                for token in numerical_score:
                    if token.ent_type_ == 'CARDINAL':
                        score = token.text
                break  # Stop looping over entities once CGPA label is found
        if found_cgpa:
            break  # Stop looping over lines once CGPA label is found

    return score



def generate_sop(template_text, res_text,programme,user_name,university):
    
    
    cgpa_score = retrieve_cgpa_score(res_text)
    print("from the terminal : ", cgpa_score)
    
    if cgpa_score != 'unknown' and float(cgpa_score) >= 7: 
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
            messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": "1st Paragraph:  introduction about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": """2nd Paragraph: Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""3rd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            
            """},
            {"role": "user", "content": f"""4th Paragraph: Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study        
            """},
            {"role": "user", "content": """5th Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            """},
            {"role": "user", "content": "6th Paragraph: Describe your future career perspectives and aspirations post-study. Show your intentions and reasons to stay and work in Germany"},
            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    
    
    elif cgpa_score != 'unknown' and float(cgpa_score) < 7:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
         messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": "1st Paragraph:  introduction about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": """2nd Paragraph: Discuss the candidate's academic and professional background :
            DON'T mention his CGPA score PLEASE.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""3rd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            
            """},
            {"role": "user", "content": f"""4th Paragraph: Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study        
            """},
            {"role": "user", "content": """5th Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            """},
            {"role": "user", "content": "6th Paragraph: Describe your future career perspectives and aspirations post-study. Show your intentions and reasons to stay and work in Germany"},
            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}
         ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    else:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
            messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": "1st Paragraph:  introduction about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": """2nd Paragraph: Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""3rd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            
            """},
            {"role": "user", "content": f"""4th Paragraph: Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study        
            """},
            {"role": "user", "content": """5th Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            """},
            {"role": "user", "content": "6th Paragraph: Describe your future career perspectives and aspirations post-study. Show your intentions and reasons to stay and work in Germany"},
            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
        

def generate_sop1(template_text, res_text,programme,user_name,university):
    
    
    cgpa_score = retrieve_cgpa_score(res_text)
    print("from the terminal : ", cgpa_score)
    
    if cgpa_score != 'unknown' and float(cgpa_score) >= 7: 
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": "1st Paragraph:  introduction about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": """2nd Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": f"""3rd Paragraph: 
            Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""4th Paragraph: Describe your future career perspectives and aspirations post-study. Show your intentions and reasons to stay and work in Germany"
                    
            """},
            {"role": "user", "content": f"""5th Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            """},
            {"role": "user", "content": """6th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Use simple words. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    
    
    elif cgpa_score != 'unknown' and float(cgpa_score) < 7:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
          messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": "1st Paragraph:  introduction about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": """2nd Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": f"""3rd Paragraph: 
            Discuss the candidate's academic and professional background :
            DON'T mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""4th Paragraph: Describe your future career perspectives and aspirations post-study. Show your intentions and reasons to stay and work in Germany"
                    
            """},
            {"role": "user", "content": f"""5th Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            """},
            {"role": "user", "content": """6th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    else:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": "1st Paragraph:  introduction about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": """2nd Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": f"""3rd Paragraph: 
            Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""4th Paragraph: Describe your future career perspectives and aspirations post-study. Show your intentions and reasons to stay and work in Germany"
                    
            """},
            {"role": "user", "content": f"""5th Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            """},
            {"role": "user", "content": """6th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Use simple words. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
        

def generate_sop2(template_text, res_text,programme,user_name,university):
    
    
    cgpa_score = retrieve_cgpa_score(res_text)
    print("from the terminal : ", cgpa_score)
    
    if cgpa_score != 'unknown' and float(cgpa_score) >= 7: 
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": f"1st Paragraph:  First sentence : Catching phrase that shows how important is studying the programme {programme} you apply for. Then introduce about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": f"""2nd Paragraph: 
            Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""3rd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            """},
            {"role": "user", "content": """4th Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": """5th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    
    
    elif cgpa_score != 'unknown' and float(cgpa_score) < 7:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": f"1st Paragraph:  First sentence : Catching phrase that shows how important is studying the programme {programme} you apply for. Then introduce about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": f"""2nd Paragraph: 
            Discuss the candidate's academic and professional background :
            DON'T mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""3rd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            """},
            {"role": "user", "content": """4th Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": """5th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    else:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": f"1st Paragraph:  First sentence : Catching phrase that shows how important is studying the programme {programme} you apply for. Then introduce about your name, background, and programme and university you are applying and also mention the reason within one line why to choose this programme"},
            {"role": "user", "content": f"""2nd Paragraph: 
            Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": f"""3rd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following : :
                Show a Strong motivation by evoking a storytelling from your past that led you to want to study this programme
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Explain Why this programme can help you build your career and prepare you for your future  
            """},
            {"role": "user", "content": """4th Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": """5th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
        

def generate_sop3(template_text, res_text,programme,user_name,university):
    
    
    cgpa_score = retrieve_cgpa_score(res_text)
    print("from the terminal : ", cgpa_score)
    
    if cgpa_score != 'unknown' and float(cgpa_score) >= 7: 
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": f"""1st Paragraph: Introduce about your name, background, and programme and university you are applying.
                                                           Then, in 4 lines you should explain how the programme you apply for is important in society and for companies
                                                           Mention that for the above 4 lines you want to study the programme to contribute to this dynamic
             """},
            {"role": "user", "content": f"""2nd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following :
                Tell what skills the programme will give you by mentionning which modules proposed you want to study
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Explain how this programme will help you become something you want to become in the future
            """},
            {"role": "user", "content": """3rd Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": """4th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": f"""5th Paragraph: 
            Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    
    
    elif cgpa_score != 'unknown' and float(cgpa_score) < 7:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": f"""1st Paragraph: Introduce about your name, background, and programme and university you are applying.
                                                           Then, in 4 lines you should explain how the programme you apply for is important in society and for companies
                                                           Mention that for the above 4 lines you want to study the programme to contribute to this dynamic
             """},
            {"role": "user", "content": f"""2nd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following :
                Tell what skills the programme will give you by mentionning which modules proposed you want to study
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Explain how this programme will help you become something you want to become in the future
            """},
            {"role": "user", "content": """3rd Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": """4th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": f"""5th Paragraph: 
            Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    else:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
           messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            # {"role": "user", "content": f"Description of the university: {university_desc}"},
            # {"role": "user", "content": f"Content of the study programme: {programme_content}"},
            {"role": "user", "content": f"""1st Paragraph: Introduce about your name, background, and programme and university you are applying.
                                                           Then, in 4 lines you should explain how the programme you apply for is important in society and for companies
                                                           Mention that for the above 4 lines you want to study the programme to contribute to this dynamic
             """},
            {"role": "user", "content": f"""2nd Paragraph: Based on the information from Internet about the programme and based on the resume : {res_text}, do the following :
                Tell what skills the programme will give you by mentionning which modules proposed you want to study
                Explain also why you are STRONGLY MOTIVATED to pursue this programme by relating to your previous experience
                Explain how this programme will help you become something you want to become in the future
            """},
            {"role": "user", "content": """3rd Paragraph: Explain why you chose to study in Germany :
            You should mention :
            Intention to stay in Germany because the field education comparative is far better to others destinations,
            Good exposure, diversity and culture 
            Mention examples of cooperations between indian and germany linked to your area of study and how you see yourself participating in it in the future ! Search for last news information you have
            
            """},
            {"role": "user", "content": """4th Paragraph:  Based on the information from Internet, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of research centers linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
                
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                    names of companies in the city of the same field of your study"""},
            {"role": "user", "content": f"""5th Paragraph: 
            Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list
            Mention all his certificates
            """},
            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
        
     

def generate_sop4(template_text, res_text,programme,user_name,university):
    
    
    cgpa_score = retrieve_cgpa_score(res_text)
    print("from the terminal : ", cgpa_score)
    
    if cgpa_score != 'unknown' and float(cgpa_score) >= 7: 
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            {"role": "user", "content": f"based on the provided details, following the specific structure and style given by the template {template_text}. You should reproduce the exact same structure of statement of purpose"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    
    
    elif cgpa_score != 'unknown' and float(cgpa_score) < 7:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            {"role": "user", "content": f"based on the provided details, following the specific structure and style given by the template {template_text}. You should reproduce the exact same structure of statement of purpose"},
            {"role": "user", "content": "DO NOT mention the Bachelor CGPA score"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    else:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
           messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"Candidate's name: {user_name}"},
            {"role": "user", "content": f"University name: {university}"},
            {"role": "user", "content": f"based on the provided details, following the specific structure and style given by the template {template_text}. You should reproduce the exact same structure of statement of purpose"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}
        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
       
       




def create_word_document(phrase, font_name, font_size):
    doc = Document()
    doc.add_paragraph(phrase)

    # Set font properties
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            
        # Set paragraph alignment to justified
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    return doc

def save_doc_to_buffer(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    
    
def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def generate_random_templates(directory_path):
    # List all files in the directory
    files = os.listdir(directory_path)
    # Filter out the PDF files
    pdf_files = [file for file in files if file.endswith('.pdf')]
    
    # Check if there are any PDF files
    if not pdf_files:
        print("No PDF files found in the directory.")
        return
    
    # Randomly select a PDF file
    selected_pdf = random.choice(pdf_files)
    print(f"Selected PDF: {selected_pdf}")
    
    # Full path to the selected PDF file
    pdf_path = os.path.join(directory_path, selected_pdf)
    return read_pdf(pdf_path)
    

def read_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)




st.markdown("""
# 📝 AI-Powered SOP Generator

Generate a sop letter : JUST READ THE INSTRUCTIONS
"""
)

# radio for upload or copy paste option         
res_format = st.radio(
    "Upload or paste the applicant's resume/key experience",
    ('Upload', 'Paste'))

if res_format == 'Upload':
    # upload_resume
    res_file = st.file_uploader('📁 Upload your resume in pdf format')
    if res_file is not None and res_file.name.endswith('.pdf'):
        pdf_reader = PdfReader(res_file)

        # Collect text from pdf
        res_text = ""
        for page in pdf_reader.pages:
            res_text += page.extract_text()
    elif res_file is not None and res_file.name.endswith('.docx'):
        st.error('sorry you should submit pdf format for the resume')

else:
        # use the pasted contents instead
        res_text = st.text_input('Pasted resume elements')
    
 
 
st.subheader("If you have a template in mind you can submit it 'Upload'.  Otherwise tick 'Let the software generate the template' and it will be automatically generated ")
# radio for upload or copy paste option         
template_format = st.radio(
    "Do you want to upload or paste the template",
    ('Upload', 'Paste', 'let the software generate the template'))

if template_format == 'Upload':     
        # upload_resume
    template_file = st.file_uploader('📁 Upload your template in pdf format')
    if template_file:
        pdf_reader = PdfReader(template_file)

        # Collect text from pdf
        template_text = ""
        for page in pdf_reader.pages:
            template_text += page.extract_text()
elif template_format == 'Paste':
    # use the pasted contents instead
    template_text = st.text_input('Pasted template elements')
else:
    template_text = generate_random_templates('templates')
            
            

with st.form('input_form'):
    # other inputs
    programme = st.text_input('Programme name')
    user_name = st.text_input('Applicant name')
    university = st.text_input('University name')
    # programme_content = st.text_input('Programme content')
    # university_desc = st.text_input('University Description')
    ai_temp = st.number_input('AI Temperature (0.0-1.0) Input how creative the API can be',value=.6)

    # submit button
    submitted = st.form_submit_button("Generate the SOP")

# if the form is submitted run the openai completion   
if submitted:
    random_number = random.randint(0, 4)
    print(random_number)
    if random_number == 0:
        response = generate_sop(template_text, res_text,programme,user_name,university)
    elif random_number == 1:
        response = generate_sop1( generate_random_templates('templates1'), res_text,programme,user_name,university)
    elif random_number == 2:
        response = generate_sop2( generate_random_templates('templates2'), res_text,programme,user_name,university)
    elif random_number == 3:
        response = generate_sop3( generate_random_templates('templates3'), res_text,programme,user_name,university)
    else:
        response = generate_sop4( generate_random_templates('templates4'), res_text,programme,user_name,university)
    
    
    # response = generate_sop4( generate_random_templates('templates4'), res_text,programme,user_name,university)
    
    doc_download1 = create_word_document(response, 'Arial', 11)
    st.download_button(
            label="Download SOP",
            data=save_doc_to_buffer(doc_download1),
            file_name=f"{res_file.name}_{university}_{programme}_SOP.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


    

