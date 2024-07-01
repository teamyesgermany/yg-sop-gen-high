import streamlit as st
import io
import os
import openai as ai
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import spacy
import json
import random
import tiktoken

import requests
from bs4 import BeautifulSoup





# Your Diffbot API token

DIFFBOT_API_URL = 'https://api.diffbot.com/v3/article'



load_dotenv()  # take environment variables from .env.


DIFFBOT_API_TOKEN = os.getenv("DIFFBOT_API_TOKEN")
ai.api_key = os.getenv("OPENAI_API_KEY")

nlp = spacy.blank("en")  # Create a blank English model
#^ Load the SpaCy English model
nlp_e = spacy.load("en_core_web_sm")






def extract_urls(query):

    # Perform the search query
    # query = "Hochschule Schmalkalden • (University) • Schmalkalden Master in Finance"
    
    
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')

        # Find the first search result link
        links = []
        results = soup.find_all('div', {'class': 'yuRUbf'})
        for result in results[:2]:
            links.append(result.find('a')['href'])
            print(result.find('a')['href'])
        else:
            print("No results found.")
    else:
        print("Failed to retrieve the website content.")
        

    # Using slicing to access the first two elements
    if len(links) >= 2:
        sliced_links = links[0:2]
        return sliced_links[0], sliced_links[1]
    # # Print each element individually
    # for i, link in enumerate(sliced_links):
    #     print(f"Link {i + 1}: {link}")



def extract_urls1(query):

    # Perform the search query
    # query = "Hochschule Schmalkalden • (University) • Schmalkalden Master in Finance"
    
    
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')

        # Find the first search result link
        links = []
        result = soup.find('div', {'class': 'yuRUbf'})
        if result:
            return result.find('a')['href']
        else:
            print("No results found.")
    else:
        print("Failed to retrieve the website content.")
        

def extract3urls(query):
      # Perform the search query
    # query = "Hochschule Schmalkalden • (University) • Schmalkalden Master in Finance"
    
    
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')

        # Find the first search result link
        links = []
        results = soup.find_all('div', {'class': 'yuRUbf'})
        for result in results[:10]:
            links.append(result.find('a')['href'])
            print(result.find('a')['href'])
        else:
            print("No results found.")
    else:
        print("Failed to retrieve the website content.")
        

    # Using slicing to access the first two elements
    if len(links) >= 10:
        sliced_links = links[0:10]
        caca = list(range(10))
        random_numbers = random.sample(caca, 3)
        print('first link :', sliced_links[random_numbers[0]], ' second link : ', sliced_links[random_numbers[1]],  "third link :", sliced_links[random_numbers[2]])
        
        return sliced_links[random_numbers[0]],  sliced_links[random_numbers[1]],  sliced_links[random_numbers[2]]
    # # Print each element individually
    # for i, link in enumerate(sliced_links):
    #     print(f"Link {i + 1}: {link}")

# Function to get data from Diffbot API
def get_data_from_diffbot(url):
    params = {
        'token': DIFFBOT_API_TOKEN,
        'url': url,
    }
    response = requests.get(DIFFBOT_API_URL, params=params)
    if response.status_code == 200:
        return response.json()
    else:
        return None


# Function to retrieve all fragment links
def get_fragment_links(base_url):
    response = requests.get(base_url)
    soup = BeautifulSoup(response.content, 'html.parser')
    fragment_links = []

    # Find all links that point to fragments within the page
    for a in soup.find_all('a', href=True):
        href = a['href']
        if href.startswith('#') and len(href) > 1:
            fragment_links.append(href)

    return fragment_links


# Collect data from each fragment
def collect_data(base_url, fragment_links):
    all_data = {}
    for fragment in fragment_links:
        full_url = f'{base_url}{fragment}'
        data = get_data_from_diffbot(full_url)
        if data:
            all_data[fragment] = data
        else:
            print(f'Failed to retrieve data from {full_url}')
    return all_data


# Main function
def return_data_withfragments(query):
    jaja = extract_urls(query)
    if jaja is not None:
        url1, url2 = jaja
    # Base URL
        bases_urls =  [url1, url2]
        programme_content = []
        for base_url in bases_urls:
            fragment_links = get_fragment_links(base_url)
            if fragment_links:
                all_data = collect_data(base_url, fragment_links)
                for fragment, data in all_data.items():
                    print(f'Data from {fragment}:')
                    print(data)
                    programme_content.append(json.dumps(data))
            else:
                print('No fragment links found.')
        return (" ".join(programme_content))
    else:
        return None
    
   
# Main function
def return_data3(query):
    if extract3urls(query) is not None:
        url1, url2, url3 = extract3urls(query)
    # Base URL
        bases_urls =  [url1, url2, url3]
        whygermany_content = []
        for base_url in bases_urls:
                # Complete API URL with your token and the specific page
            api_url = f"{DIFFBOT_API_URL}?token={DIFFBOT_API_TOKEN}&url={base_url}"
            
            # Making the HTTP GET request to Diffbot Custom API
            response = requests.get(api_url)

            # Checking the response
            if response.status_code == 200:
                # Printing the structured data extracted by your custom model
                print(response.json())
                whygermany_content.append(json.dumps(response.json()))
            else:
                print("Failed to retrieve data:", response.status_code)
                    
        return (" ".join(whygermany_content))
    else:
        return None
    

# Main function
def return_data1(query):
    if extract_urls1(query) is not None:
        url1 = extract_urls1(query)
        
        # Complete API URL with your token and the specific page
        api_url = f"{DIFFBOT_API_URL}?token={DIFFBOT_API_TOKEN}&url={url1}"
        
        # Making the HTTP GET request to Diffbot Custom API
        response = requests.get(api_url)

        # Checking the response
        if response.status_code == 200:
            # Printing the structured data extracted by your custom model
            print(response.json())
            return response.json()
        else:
            print("Failed to retrieve data:", response.status_code)
        
    else:
        return None




def retrieve_cgpa_score(content_resume):
    
    # Check if the entity_ruler already exists in the pipeline
    if 'entity_ruler' in nlp.pipe_names:
        ruler = nlp.get_pipe('entity_ruler')
    else:
        ruler = nlp.add_pipe("entity_ruler")

    patterns = [
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.6 CGPA"
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "CGPA 9.6" and "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]},  # Matches "CGPA - 9.26" or "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "9.5 CGPA" or "9.5 CGPA -" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.5 - CGPA" or "9.5 . CGPA" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["-", "–"]}}, {"IS_SPACE": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]}, # Matches "CGPA - 9.26" or "CGPA -9.5" (with optional punctuation and space)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\w*\\d+(\\.\\d+)?"}}, {"LOWER": {"REGEX": "(cgpa|sgpi\\w*)"}}]},
        {"label": "CGPA", "pattern": [{"LOWER": {"REGEX": "(\\w*cgpa|sgpi)"}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?\w*"}}]}
    ]

    ruler.add_patterns(patterns)

    lines = content_resume.split('\n')
    score = "unknown"
    for index, line in enumerate(lines):
        doc = nlp(line)
        found_cgpa = False
        for ent in doc.ents:
            if ent.label_ == 'CGPA':
                found_cgpa = True
                print(ent.text)
                numerical_score = nlp_e(ent.text)
                for token in numerical_score:
                    if token.ent_type_ == 'CARDINAL':
                        score = token.text
                break  # Stop looping over entities once CGPA label is found
        if found_cgpa:
            break  # Stop looping over lines once CGPA label is found

    return score


def generate_sop4(template_text, res_text,programme, university, programme_content, university_description, facilities, research_institutes, ranking, location, international_students, modules, career, partnerships, cooperation_india_germany, germany):
    
    
    enc = tiktoken.encoding_for_model("gpt-4o")
     # Concatenate all text inputs to calculate total tokens
    all_texts = [programme, university, programme_content, university_description, facilities, research_institutes, ranking, location, international_students, modules, career, partnerships, cooperation_india_germany, germany]
    all_texts_tokens = sum(len(enc.encode(json.dumps(text))) for text in all_texts)
    print("total number of tokens is", all_texts_tokens)
    
    
        # Define the token limit for the entire prompt
    if all_texts_tokens > 128000:
        germany = ""
        cooperation_india_germany = ""
    print("total number of tokens is", all_texts_tokens)
        
    
    cgpa_score = retrieve_cgpa_score(res_text)
    print("from the terminal : ", cgpa_score)
    
    if cgpa_score != 'unknown' and float(cgpa_score) >= 7: 
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        
  messages = [
    {"role": "user", "content": f"Study programme name: {programme}"},
    {"role": "user", "content": f"University name: {university}"},
    {"role": "user", "content": f"You have to answer the following questions:"},
    {"role": "user", "content": f"Imagine yourself as an applicant who wants to integrate the university {university} for the programme {programme}."},
    {"role": "user", "content": f"For each of the following questions, I will tell you, on which information you will have to be based:"},
    
    {"role": "user", "content": f"Question 1: Provide detailed examples and facts about the expertise and reputation of the faculty for the {programme} at {university}. Include specific names, projects, publications, awards, and collaborations that illustrate the faculty's strengths. Retrieve information from the university website, Google News, and the internet in general."},
     
    {"role": "user", "content": f"Question 2: What unique facilities does {university} offer that are not commonly available at other universities? Base your response on {facilities} and {university_description}"},
    
    {"role": "user", "content": f"Question 3: List all the research projects and research centers associated with the {programme} at {university}.Base your response on {research_institutes} and {university_description}"},
    
    {"role": "user", "content": f"Question 4: What is the ranking of {university} according to reliable ranking sources? Provide the exact rank in Germany and in the world. Base your response on {ranking} and {university_description}"},
    
    {"role": "user", "content": f"Question 5: How does the location of {university} offer something particular and better for students compared to other universities? Base your response on {location} and {university_description}"},
    
    {"role": "user", "content": f"Question 6: How many international students are at {university}, and how many come from India? Explain why it is beneficial to have a campus with international students.Base your response on {international_students} and {university_description}"},
    
    {"role": "user", "content": f"Question 7: What is the potential of the {programme} for you? What new things will you learn from the modules? Use the following programme contents for reference: {programme_content} and {modules}."},
    
    {"role": "user", "content": f"Question 8: How does the {programme} prepare you for your professional career? Base your response on {career} and {programme_content}"},
    
    {"role": "user", "content": f"Question 9: What industry partnerships does {university} have that can be interesting for you? Base your response on {partnerships}"},
    
    {"role": "user", "content": f"Question 10: Which technical skills will you learn from the {programme}? Use the following programme contents for reference: {programme_content} and {modules}."},
    
    {"role": "user", "content": f"Question 11: What elements from your resume ({res_text}) motivated you to pursue the {programme}?"},
    
    {"role": "user", "content": f"Question 12: Invent an imaginative anecdote based on your resume ({res_text}) that gave you the interest to study the {programme}."},
    
    {"role": "user", "content": f"Question 13: Talk about cooperations between India and Germany that imagine yourself participating in, in the future. Your response should be based on {cooperation_india_germany}"},

    {"role": "user", "content": f"""Question 14: What factors make Germany an ideal study destination for you? Discuss the following points with examples relevant to this programme : (Your responses should be based on {germany}):
        1. Quality Education: Explain why Germany's high-quality education system is appealing to you. 
        2. Program Variety: Discuss how the variety of programs available in Germany aligns with your academic interests and goals.
        3. Economic Opportunities: Discuss how the opportunities available in Germany align with your career goals.
        4. Research Opportunities: Highlight the research opportunities in your field of study that attract you to Germany.
        5. Employment Opportunities: Discuss the employment opportunities available in Germany that align with your career goals.
        6. Mention examples of cooperation between India and Germany relevant to this programme.
        7. Cultural Experience: Describe how experiencing German culture and language will enrich your personal and academic growth.
        8. Financial Support: Explain how scholarships and financial support options in Germany will help alleviate your financial burden.
        9. Standard of Living: Discuss the high standard of living and quality of life that Germany offers to students.
        10.Location and Industry Links: Explain the strategic advantages of Germany's location and its strong ties to industries relevant to your field.
        11.Tuition Fees: Discuss the affordability of education in Germany due to its low tuition fees.
        12.Post-Graduation Opportunities: Highlight the post-graduation opportunities, including visa extensions, that make Germany a favorable choice for your future career prospects.
        """},
    
    {"role": "user", "content": "Provide concrete examples for Questions 1 to 14. For each question , give 10 lines of response minimum. I want DETAILED and RELEVANT information that goes beyond a simple sentence."},
    {"role": "user", "content": "The responses should be human like."},
    {"role": "user", "content": "I repeat that the responses should be human like that means very SIMPLE  , responses understood by 18 years old people."},


        ]

        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    
    
    elif cgpa_score != 'unknown' and float(cgpa_score) < 7:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"University name: {university}"},
            {"role": "user", "content": f"based on the provided details, following the specific structure and style given by the template {template_text}. You should reproduce the exact same structure of statement of purpose"},
            {"role": "user", "content": "DO NOT mention the Bachelor CGPA score"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    else:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
           messages=[
            {"role": "user", "content": "Generate a statement of purpose based on the provided details, following a specific structure and style."},
            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Resume information: {res_text}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"University name: {university}"},
            {"role": "user", "content": f"based on the provided details, following the specific structure and style given by the template {template_text}. You should reproduce the exact same structure of statement of purpose"},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on', 'delve', 'renowned'"}
        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
       
       
def create_word_document(phrase, font_name, font_size):
    doc = Document()
    doc.add_paragraph(phrase)

    # Set font properties
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            
        # Set paragraph alignment to justified
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    return doc


def save_doc_to_buffer(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    
    
def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def generate_random_templates(directory_path):
    # List all files in the directory
    files = os.listdir(directory_path)
    # Filter out the PDF files
    pdf_files = [file for file in files if file.endswith('.pdf')]
    
    # Check if there are any PDF files
    if not pdf_files:
        print("No PDF files found in the directory.")
        return
    
    # Randomly select a PDF file
    selected_pdf = random.choice(pdf_files)
    print(f"Selected PDF: {selected_pdf}")
    
    # Full path to the selected PDF file
    pdf_path = os.path.join(directory_path, selected_pdf)
    return read_pdf(pdf_path)
    

def read_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)




st.markdown("""
# 📝 AI-Powered SOP Generator

Generate a sop letter : JUST READ THE INSTRUCTIONS
"""
)

# radio for upload or copy paste option         
res_format = st.radio(
    "Upload or paste the applicant's resume/key experience",
    ('Upload', 'Paste'))

if res_format == 'Upload':
    # upload_resume
    res_file = st.file_uploader('📁 Upload your resume in pdf format')
    if res_file is not None and res_file.name.endswith('.pdf'):
        pdf_reader = PdfReader(res_file)

        # Collect text from pdf
        res_text = ""
        for page in pdf_reader.pages:
            res_text += page.extract_text()
    elif res_file is not None and res_file.name.endswith('.docx'):
        st.error('sorry you should submit pdf format for the resume')

else:
        # use the pasted contents instead
        res_text = st.text_input('Pasted resume elements')
    
 
 
st.subheader("If you have a template in mind you can submit it 'Upload'.  Otherwise tick 'Let the software generate the template' and it will be automatically generated ")
# radio for upload or copy paste option         
template_format = st.radio(
    "Do you want to upload or paste the template",
    ('Upload', 'Paste', 'let the software generate the template'))

if template_format == 'Upload':     
        # upload_resume
    template_file = st.file_uploader('📁 Upload your template in pdf format')
    if template_file:
        pdf_reader = PdfReader(template_file)

        # Collect text from pdf
        template_text = ""
        for page in pdf_reader.pages:
            template_text += page.extract_text()
elif template_format == 'Paste':
    # use the pasted contents instead
    template_text = st.text_input('Pasted template elements')
else:
    template_text = generate_random_templates('templates')
            
            

with st.form('input_form'):
    # other inputs
    programme = st.text_input('Programme name')
    university = st.text_input('University name')
    ai_temp = st.number_input('AI Temperature (0.0-1.0) Input how creative the API can be',value=.6)
    

    
    # # submit button
    submitted = st.form_submit_button("Generate the SOP")

# if the form is submitted run the openai completion   
if submitted:
    research = university + " " + programme
    programme_content = return_data_withfragments(research)  #with fragments
    university_description = return_data1(f"{university} wikipedia")
    facilities = return_data1(f"What unique facilities does {university} offer that are not commonly available at other universities")
    research_institutes = return_data1(f"{university},school of {programme},research institutes")
    ranking = return_data1(f"ranking {university}")
    location = return_data1(f"How does the location of {university} offers something particular and better for students compared to other universities")
    international_students = return_data1(f"How many international students are at {university}, and how many come from India? Explain why it is beneficial to have a campus with international students.")
    modules = return_data1(f"{programme}, {university}, modules") #with fragments
    career = return_data1(f"How does the {programme} prepare you for your professional career?")
    partnerships = return_data1(f"partnerships {university}")
    cooperation_india_germany = return_data1(f"bilateral cooperation between india and germany , {programme}")
    germany = return_data1("What factors make Germany an ideal study destination for you?")
    

    response = generate_sop4( generate_random_templates('templates4'), res_text,programme,university, programme_content, university_description, facilities, research_institutes, ranking, location, international_students, modules, career, partnerships, cooperation_india_germany, germany)
    
    doc_download1 = create_word_document(response, 'Arial', 11)
    st.download_button(
            label="Download SOP",
            data=save_doc_to_buffer(doc_download1),
            file_name=f"{res_file.name}_{university}_{programme}_SOP.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


    

